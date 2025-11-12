from fastapi import APIRouter, HTTPException, Depends, status
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from pydantic import BaseModel
from typing import List, Dict, Any, Optional, Union
import logging
import os
import re
import unicodedata
import io
from datetime import datetime
import tempfile

# Import report generation libraries
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor, black, blue
from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

from database_simple import get_db, User, Document, Conversation, Message, ExportHistory
from routers.auth import get_current_user
from security.csrf import verify_csrf as csrf_protect

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

router = APIRouter(tags=["exports"])

# Pydantic models
class ExportRequest(BaseModel):
    conversation_id: str
    export_type: str  # pdf, docx, txt
    include_citations: bool = True
    title: Optional[str] = None

class ExportResponse(BaseModel):
    success: bool
    filename: str
    file_size: int
    download_url: str
    export_id: str
    error: Optional[str] = None

class ExportHistoryResponse(BaseModel):
    id: str
    document_title: str
    export_type: str
    filename: str
    file_size: int
    created_at: str

class ArtifactExportRequest(BaseModel):
    document_id: str
    artifact_type: str  # e.g., 'study_guide'
    export_type: str    # pdf | docx | txt
    source_format: str  # markdown | json
    title: Optional[str] = None
    content: Union[str, Dict[str, Any]]

# Ensure exports directory exists
# Exports directory (absolute path safer for file responses)
EXPORTS_DIR = "/home/angel/DocAI/apps/api/exports"
os.makedirs(EXPORTS_DIR, exist_ok=True)


def _safe_base_filename(name: str, desired_ext: str) -> str:
    """Sanitize a filename base, strip known extensions, remove diacritics and unsafe chars."""
    if not name:
        name = "export"
    # Strip known ext if present
    lower = name.lower()
    for ext in (".pdf", ".docx", ".txt", ".pptx"):
        if lower.endswith(ext):
            name = name[: -len(ext)]
            break
    # Remove diacritics and non-ascii
    name_ascii = unicodedata.normalize("NFKD", name).encode("ascii", "ignore").decode()
    # Replace spaces with underscores
    name_ascii = re.sub(r"\s+", "_", name_ascii)
    # Keep only safe chars
    name_ascii = re.sub(r"[^A-Za-z0-9._-]", "_", name_ascii)
    name_ascii = name_ascii.strip("._-") or "export"
    # Trim length and ensure ext
    name_ascii = name_ascii[:120]
    return f"{name_ascii}.{desired_ext}"

def clean_text_for_export(text: str) -> str:
    """Clean text for export by removing markdown and extra whitespace"""
    if not text:
        return ""
    
    # Remove common markdown patterns
    cleaned = text.replace("**", "").replace("*", "")
    cleaned = cleaned.replace("##", "").replace("#", "")
    
    # Clean up whitespace
    lines = [line.strip() for line in cleaned.split('\n') if line.strip()]
    return '\n'.join(lines)

def generate_pdf_export(
    conversation_data: Dict[str, Any],
    include_citations: bool = True,
    custom_title: Optional[str] = None
) -> str:
    """Generate PDF export of conversation"""
    try:
        # Create temporary file
        temp_file = tempfile.NamedTemporaryFile(
            delete=False, 
            suffix='.pdf',
            dir=EXPORTS_DIR
        )
        temp_file.close()
        
        # Create PDF document
        doc = SimpleDocTemplate(
            temp_file.name,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # Get styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            textColor=HexColor('#1f2937'),
            alignment=1  # Center
        )
        
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Heading2'],
            fontSize=16,
            spaceAfter=20,
            textColor=HexColor('#4b5563')
        )
        
        user_style = ParagraphStyle(
            'UserMessage',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            leftIndent=0,
            textColor=HexColor('#1f2937'),
            backColor=HexColor('#f9fafb')
        )
        
        assistant_style = ParagraphStyle(
            'AssistantMessage',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            leftIndent=0,
            textColor=HexColor('#1f2937')
        )
        
        citation_style = ParagraphStyle(
            'Citation',
            parent=styles['Normal'],
            fontSize=9,
            spaceAfter=6,
            leftIndent=20,
            textColor=HexColor('#6b7280'),
            backColor=HexColor('#f3f4f6')
        )
        
        # Build document content
        story = []
        
        # Title
        title = custom_title or f"Conversación: {conversation_data['document_title']}"
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 12))
        
        # Metadata
        export_date = datetime.now().strftime('%d/%m/%Y %H:%M')
        story.append(Paragraph(f"<b>Documento:</b> {conversation_data['document_title']}", subtitle_style))
        story.append(Paragraph(f"<b>Fecha de exportación:</b> {export_date}", styles['Normal']))
        story.append(Paragraph(f"<b>Total de mensajes:</b> {len(conversation_data['messages'])}", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Messages
        for i, message in enumerate(conversation_data['messages']):
            # Message header
            role_text = "Usuario" if message['role'] == 'user' else "DocAI Assistant"
            timestamp = datetime.fromisoformat(message['created_at'].replace('Z', '+00:00')).strftime('%H:%M')
            
            header = f"<b>{role_text}</b> - {timestamp}"
            style = user_style if message['role'] == 'user' else assistant_style
            
            story.append(Paragraph(header, subtitle_style))
            
            # Message content
            clean_content = clean_text_for_export(message['content'])
            story.append(Paragraph(clean_content, style))
            
            # Citations
            if include_citations and message.get('citations') and len(message['citations']) > 0:
                story.append(Spacer(1, 6))
                story.append(Paragraph("<b>Citas del documento:</b>", citation_style))
                
                for citation in message['citations']:
                    citation_text = f"Página {citation['page']}: \"{citation['snippet']}\""
                    story.append(Paragraph(citation_text, citation_style))
            
            story.append(Spacer(1, 15))
            
            # Add page break after every 10 messages for readability
            if (i + 1) % 10 == 0 and i < len(conversation_data['messages']) - 1:
                story.append(PageBreak())
        
        # Build PDF
        doc.build(story)
        
        return temp_file.name
        
    except Exception as e:
        logger.error(f"Error generating PDF: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to generate PDF: {str(e)}")

def generate_study_guide_pdf(
    guide_data: Union[str, Dict[str, Any]],
    source_format: str = "markdown",
    custom_title: Optional[str] = None,
    document_title: Optional[str] = None,
) -> str:
    """Generate PDF for a study guide artifact.
    - If source_format == 'json', expects structured dict with fields.
    - If source_format == 'markdown', treats content as plain markdown text (no full markdown rendering).
    """
    try:
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf', dir=EXPORTS_DIR)
        temp_file.close()

        doc = SimpleDocTemplate(
            temp_file.name,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle('GuideTitle', parent=styles['Heading1'], fontSize=22, spaceAfter=14, textColor=HexColor('#1f2937'), alignment=1)
        h2_style = ParagraphStyle('GuideH2', parent=styles['Heading2'], fontSize=15, spaceBefore=10, spaceAfter=8, textColor=HexColor('#1f2937'))
        body_style = ParagraphStyle('GuideBody', parent=styles['Normal'], fontSize=11, spaceAfter=8, textColor=HexColor('#1f2937'))
        meta_style = ParagraphStyle('GuideMeta', parent=styles['Normal'], fontSize=9, textColor=HexColor('#6b7280'))
        bullet_style = ParagraphStyle('GuideBullet', parent=styles['Normal'], fontSize=11, leftIndent=14, spaceAfter=4, textColor=HexColor('#1f2937'))

        story = []

        title = custom_title or f"Guía de Estudio"
        if document_title:
            title += f" – {document_title}"
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 6))

        export_date = datetime.now().strftime('%d/%m/%Y %H:%M')
        if document_title:
            story.append(Paragraph(f"Documento: {document_title}", meta_style))
        story.append(Paragraph(f"Fecha de exportación: {export_date}", meta_style))
        story.append(Spacer(1, 12))

        if source_format == 'markdown' and isinstance(guide_data, str):
            for line in guide_data.split('\n'):
                text = line.strip()
                if not text:
                    story.append(Spacer(1, 6))
                    continue
                if text.startswith('## '):
                    story.append(Paragraph(text[3:], h2_style))
                elif text.startswith('- '):
                    story.append(Paragraph(text, bullet_style))
                else:
                    story.append(Paragraph(text, body_style))
            doc.build(story)
            return temp_file.name

        # JSON structured
        if isinstance(guide_data, dict):
            gd = guide_data
            if gd.get('objectives'):
                story.append(Paragraph('Objetivos', h2_style))
                for o in gd['objectives']:
                    story.append(Paragraph(f"- {o}", bullet_style))
                story.append(Spacer(1, 6))
            if gd.get('key_concepts'):
                story.append(Paragraph('Conceptos clave', h2_style))
                for kc in gd['key_concepts']:
                    term = kc.get('term', 'Concepto')
                    definition = kc.get('definition', '')
                    pages = kc.get('pages', '')
                    text = f"- <b>{term}</b>: {definition} {f'[{pages}]' if pages else ''}"
                    story.append(Paragraph(text, bullet_style))
                story.append(Spacer(1, 6))
            if gd.get('sections'):
                story.append(Paragraph('Secciones', h2_style))
                for s in gd['sections']:
                    stitle = s.get('title', 'Sección')
                    pages = s.get('pages', '')
                    story.append(Paragraph(f"{stitle} {f'[{pages}]' if pages else ''}", body_style))
                    summary = s.get('summary')
                    if summary:
                        story.append(Paragraph(summary, body_style))
                    kps = s.get('key_points') or []
                    for kp in kps:
                        story.append(Paragraph(f"- {kp}", bullet_style))
                    exs = s.get('examples') or []
                    for ex in exs:
                        story.append(Paragraph(f"- {ex}", bullet_style))
                    story.append(Spacer(1, 6))
            if gd.get('checkpoints'):
                story.append(Paragraph('Checkpoints', h2_style))
                for c in gd['checkpoints']:
                    story.append(Paragraph(f"- {c}", bullet_style))
                story.append(Spacer(1, 6))
            if gd.get('review_questions'):
                story.append(Paragraph('Preguntas de repaso', h2_style))
                for q in gd['review_questions']:
                    txt = q.get('question', 'Pregunta') if isinstance(q, dict) else str(q)
                    pages = q.get('pages', '') if isinstance(q, dict) else ''
                    story.append(Paragraph(f"- {txt} {f'[{pages}]' if pages else ''}", bullet_style))
                story.append(Spacer(1, 6))
            if gd.get('estimated_time_minutes'):
                story.append(Paragraph(f"Tiempo estimado: {gd['estimated_time_minutes']} min", meta_style))

            doc.build(story)
            return temp_file.name

        # Fallback: plain text
        text = str(guide_data)
        for line in text.split('\n'):
            story.append(Paragraph(line, body_style))
        doc.build(story)
        return temp_file.name
    except Exception as e:
        logger.error(f"Error generating Study Guide PDF: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to generate Study Guide PDF: {str(e)}")

def generate_study_guide_docx(
    guide_data: Union[str, Dict[str, Any]],
    source_format: str = "markdown",
    custom_title: Optional[str] = None,
    document_title: Optional[str] = None,
) -> str:
    try:
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx', dir=EXPORTS_DIR)
        temp_file.close()
        doc = DocxDocument()
        title = custom_title or "Guía de Estudio"
        if document_title:
            title += f" – {document_title}"
        doc.add_heading(title, 0)

        if source_format == 'markdown' and isinstance(guide_data, str):
            for line in guide_data.split('\n'):
                t = line.strip()
                if not t:
                    doc.add_paragraph("")
                    continue
                if t.startswith('## '):
                    doc.add_heading(t[3:], level=2)
                elif t.startswith('- '):
                    p = doc.add_paragraph(t[2:])
                    p.style = 'List Bullet'
                else:
                    doc.add_paragraph(t)
            doc.save(temp_file.name)
            return temp_file.name

        if isinstance(guide_data, dict):
            gd = guide_data
            if gd.get('objectives'):
                doc.add_heading('Objetivos', level=2)
                for o in gd['objectives']:
                    p = doc.add_paragraph(o)
                    p.style = 'List Bullet'
            if gd.get('key_concepts'):
                doc.add_heading('Conceptos clave', level=2)
                for kc in gd['key_concepts']:
                    term = kc.get('term', 'Concepto')
                    definition = kc.get('definition', '')
                    pages = kc.get('pages', '')
                    doc.add_paragraph(f"{term}: {definition} {f'[{pages}]' if pages else ''}")
            if gd.get('sections'):
                doc.add_heading('Secciones', level=2)
                for s in gd['sections']:
                    stitle = s.get('title', 'Sección')
                    pages = s.get('pages', '')
                    doc.add_heading(f"{stitle} {f'[{pages}]' if pages else ''}", level=3)
                    summary = s.get('summary')
                    if summary:
                        doc.add_paragraph(summary)
                    for kp in s.get('key_points') or []:
                        p = doc.add_paragraph(kp)
                        p.style = 'List Bullet'
                    for ex in s.get('examples') or []:
                        p = doc.add_paragraph(ex)
                        p.style = 'List Bullet'
            if gd.get('checkpoints'):
                doc.add_heading('Checkpoints', level=2)
                for c in gd['checkpoints']:
                    p = doc.add_paragraph(c)
                    p.style = 'List Bullet'
            if gd.get('review_questions'):
                doc.add_heading('Preguntas de repaso', level=2)
                for q in gd['review_questions']:
                    txt = q.get('question', 'Pregunta') if isinstance(q, dict) else str(q)
                    pages = q.get('pages', '') if isinstance(q, dict) else ''
                    doc.add_paragraph(f"{txt} {f'[{pages}]' if pages else ''}")
            if gd.get('estimated_time_minutes'):
                doc.add_paragraph(f"Tiempo estimado: {gd['estimated_time_minutes']} min")

            doc.save(temp_file.name)
            return temp_file.name

        # Fallback plain text
        doc.add_paragraph(str(guide_data))
        doc.save(temp_file.name)
        return temp_file.name
    except Exception as e:
        logger.error(f"Error generating Study Guide DOCX: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to generate Study Guide DOCX: {str(e)}")

def generate_study_guide_txt(
    guide_data: Union[str, Dict[str, Any]],
    source_format: str = "markdown",
    custom_title: Optional[str] = None,
    document_title: Optional[str] = None,
) -> str:
    try:
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.txt', dir=EXPORTS_DIR, mode='w', encoding='utf-8')
        title = custom_title or "Guía de Estudio"
        if document_title:
            title += f" – {document_title}"
        temp_file.write(title + "\n" + ("="*max(12, len(title))) + "\n\n")
        if source_format == 'markdown' and isinstance(guide_data, str):
            temp_file.write(guide_data)
        elif isinstance(guide_data, dict):
            gd = guide_data
            if gd.get('objectives'):
                temp_file.write("Objetivos\n")
                for o in gd['objectives']:
                    temp_file.write(f"- {o}\n")
                temp_file.write("\n")
            if gd.get('key_concepts'):
                temp_file.write("Conceptos clave\n")
                for kc in gd['key_concepts']:
                    term = kc.get('term', 'Concepto')
                    definition = kc.get('definition', '')
                    pages = kc.get('pages', '')
                    temp_file.write(f"- {term}: {definition} {f'[{pages}]' if pages else ''}\n")
                temp_file.write("\n")
            if gd.get('sections'):
                temp_file.write("Secciones\n")
                for s in gd['sections']:
                    stitle = s.get('title', 'Sección')
                    pages = s.get('pages', '')
                    temp_file.write(f"{stitle} {f'[{pages}]' if pages else ''}\n")
                    if s.get('summary'):
                        temp_file.write(f"  - Síntesis: {s['summary']}\n")
                    for kp in s.get('key_points') or []:
                        temp_file.write(f"  - {kp}\n")
                    for ex in s.get('examples') or []:
                        temp_file.write(f"  - {ex}\n")
                    temp_file.write("\n")
            if gd.get('checkpoints'):
                temp_file.write("Checkpoints\n")
                for c in gd['checkpoints']:
                    temp_file.write(f"- {c}\n")
                temp_file.write("\n")
            if gd.get('review_questions'):
                temp_file.write("Preguntas de repaso\n")
                for q in gd['review_questions']:
                    txt = q.get('question', 'Pregunta') if isinstance(q, dict) else str(q)
                    pages = q.get('pages', '') if isinstance(q, dict) else ''
                    temp_file.write(f"- {txt} {f'[{pages}]' if pages else ''}\n")
                temp_file.write("\n")
            if gd.get('estimated_time_minutes'):
                temp_file.write(f"Tiempo estimado: {gd['estimated_time_minutes']} min\n")
        else:
            temp_file.write(str(guide_data))
        temp_file.close()
        return temp_file.name
    except Exception as e:
        logger.error(f"Error generating Study Guide TXT: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to generate Study Guide TXT: {str(e)}")

def generate_docx_export(
    conversation_data: Dict[str, Any],
    include_citations: bool = True,
    custom_title: Optional[str] = None
) -> str:
    """Generate Word document export of conversation"""
    try:
        # Create temporary file
        temp_file = tempfile.NamedTemporaryFile(
            delete=False, 
            suffix='.docx',
            dir=EXPORTS_DIR
        )
        temp_file.close()
        
        # Create Word document
        doc = DocxDocument()
        
        # Title
        title = custom_title or f"Conversación: {conversation_data['document_title']}"
        title_paragraph = doc.add_heading(title, 0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        export_date = datetime.now().strftime('%d/%m/%Y %H:%M')
        doc.add_paragraph(f"Documento: {conversation_data['document_title']}")
        doc.add_paragraph(f"Fecha de exportación: {export_date}")
        doc.add_paragraph(f"Total de mensajes: {len(conversation_data['messages'])}")
        doc.add_paragraph("")
        
        # Messages
        for message in conversation_data['messages']:
            # Message header
            role_text = "Usuario" if message['role'] == 'user' else "DocAI Assistant"
            timestamp = datetime.fromisoformat(message['created_at'].replace('Z', '+00:00')).strftime('%H:%M')
            
            header_paragraph = doc.add_heading(f"{role_text} - {timestamp}", level=2)
            
            # Message content
            clean_content = clean_text_for_export(message['content'])
            content_paragraph = doc.add_paragraph(clean_content)
            
            # Style based on role
            if message['role'] == 'user':
                # User message styling
                content_paragraph.style = 'Intense Quote'
            
            # Citations
            if include_citations and message.get('citations') and len(message['citations']) > 0:
                citations_header = doc.add_paragraph("Citas del documento:")
                citations_header.bold = True
                
                for citation in message['citations']:
                    citation_text = f"Página {citation['page']}: \"{citation['snippet']}\""
                    citation_paragraph = doc.add_paragraph(citation_text)
                    citation_paragraph.style = 'List Bullet'
            
            doc.add_paragraph("")  # Space between messages
        
        # Save document
        doc.save(temp_file.name)
        
        return temp_file.name
        
    except Exception as e:
        logger.error(f"Error generating DOCX: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to generate DOCX: {str(e)}")

def generate_txt_export(
    conversation_data: Dict[str, Any],
    include_citations: bool = True,
    custom_title: Optional[str] = None
) -> str:
    """Generate plain text export of conversation"""
    try:
        # Create temporary file
        temp_file = tempfile.NamedTemporaryFile(
            delete=False, 
            suffix='.txt',
            dir=EXPORTS_DIR,
            mode='w',
            encoding='utf-8'
        )
        
        # Write content
        title = custom_title or f"Conversación: {conversation_data['document_title']}"
        export_date = datetime.now().strftime('%d/%m/%Y %H:%M')
        
        temp_file.write("=" * 60 + "\n")
        temp_file.write(f"{title}\n")
        temp_file.write("=" * 60 + "\n\n")
        
        temp_file.write(f"Documento: {conversation_data['document_title']}\n")
        temp_file.write(f"Fecha de exportación: {export_date}\n")
        temp_file.write(f"Total de mensajes: {len(conversation_data['messages'])}\n\n")
        temp_file.write("-" * 60 + "\n\n")
        
        # Messages
        for message in conversation_data['messages']:
            # Message header
            role_text = "USUARIO" if message['role'] == 'user' else "DOCAI ASSISTANT"
            timestamp = datetime.fromisoformat(message['created_at'].replace('Z', '+00:00')).strftime('%H:%M')
            
            temp_file.write(f"[{role_text}] - {timestamp}\n")
            temp_file.write("-" * 40 + "\n")
            
            # Message content
            clean_content = clean_text_for_export(message['content'])
            temp_file.write(f"{clean_content}\n")
            
            # Citations
            if include_citations and message.get('citations') and len(message['citations']) > 0:
                temp_file.write("\nCitas del documento:\n")
                for citation in message['citations']:
                    temp_file.write(f"  • Página {citation['page']}: \"{citation['snippet']}\"\n")
            
            temp_file.write("\n" + "=" * 60 + "\n\n")
        
        temp_file.close()
        
        return temp_file.name
        
    except Exception as e:
        logger.error(f"Error generating TXT: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to generate TXT: {str(e)}")

@router.post("/conversations/export", response_model=ExportResponse)
async def export_conversation(
    export_request: ExportRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
    _=Depends(csrf_protect),
):
    """Export conversation in specified format"""
    try:
        # Verify conversation belongs to user
        conversation = db.query(Conversation).filter(
            Conversation.id == export_request.conversation_id,
            Conversation.user_id == current_user.id
        ).first()
        
        if not conversation:
            raise HTTPException(
                status_code=404,
                detail="Conversation not found"
            )
        
        # Get document info
        document = db.query(Document).filter(
            Document.id == conversation.document_id
        ).first()
        
        if not document:
            raise HTTPException(
                status_code=404,
                detail="Document not found"
            )
        
        # Get messages
        messages = db.query(Message).filter(
            Message.conversation_id == export_request.conversation_id
        ).order_by(Message.created_at).all()
        
        if not messages:
            raise HTTPException(
                status_code=400,
                detail="No messages found in conversation"
            )
        
        # Prepare conversation data
        conversation_data = {
            "document_title": document.title,
            "messages": [
                {
                    "role": msg.role,
                    "content": msg.content,
                    "citations": msg.citations or [],
                    "created_at": msg.created_at.isoformat()
                }
                for msg in messages
            ]
        }
        
        # Generate export based on type
        if export_request.export_type == "pdf":
            file_path = generate_pdf_export(
                conversation_data, 
                export_request.include_citations,
                export_request.title
            )
        elif export_request.export_type == "docx":
            file_path = generate_docx_export(
                conversation_data, 
                export_request.include_citations,
                export_request.title
            )
        elif export_request.export_type == "txt":
            file_path = generate_txt_export(
                conversation_data, 
                export_request.include_citations,
                export_request.title
            )
        else:
            raise HTTPException(
                status_code=400,
                detail="Invalid export type. Must be 'pdf', 'docx', or 'txt'"
            )
        
        # Get file size
        file_size = os.path.getsize(file_path)
        filename = os.path.basename(file_path)
        
        # Save export history
        export_history = ExportHistory(
            user_id=current_user.id,
            conversation_id=export_request.conversation_id,
            document_id=conversation.document_id,
            export_type=export_request.export_type,
            filename=filename,
            file_size=file_size
        )
        db.add(export_history)
        db.commit()
        
        return ExportResponse(
            success=True,
            filename=filename,
            file_size=file_size,
            download_url=f"/api/exports/download/{filename}",
            export_id=export_history.id
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Export error: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to export conversation: {str(e)}"
        )

@router.get("/download/{filename}")
async def download_export(
    filename: str,
    current_user: User = Depends(get_current_user)
):
    """Download exported file"""
    try:
        file_path = os.path.join(EXPORTS_DIR, filename)

        if not os.path.exists(file_path):
            # Debug log: list available files to help diagnose
            try:
                logger.warning(f"Download missing: {file_path}")
                logger.warning(f"Available: {', '.join(os.listdir(EXPORTS_DIR))}")
            except Exception:
                pass
            raise HTTPException(
                status_code=404,
                detail="File not found"
            )

        # Determine media type based on extension
        if filename.endswith('.pdf'):
            media_type = 'application/pdf'
        elif filename.endswith('.docx'):
            media_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif filename.endswith('.txt'):
            media_type = 'text/plain'
        else:
            media_type = 'application/octet-stream'

        return FileResponse(
            path=file_path,
            filename=filename,
            media_type=media_type
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Download error: {e}")
        raise HTTPException(
            status_code=500,
            detail="Failed to download file"
        )

@router.post("/artifacts/export", response_model=ExportResponse)
async def export_artifact(
    req: ArtifactExportRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
    _=Depends(csrf_protect),
):
    """Export arbitrary artifacts such as study guides (PDF/DOCX/TXT)."""
    try:
        # Validate document belongs to user
        document = db.query(Document).filter(
            Document.id == req.document_id,
            Document.user_id == current_user.id
        ).first()
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        artifact_type = req.artifact_type.lower()
        export_type = req.export_type.lower()
        source_format = req.source_format.lower()
        if artifact_type not in {"study_guide"}:
            raise HTTPException(status_code=400, detail="Unsupported artifact type")
        if export_type not in {"pdf", "docx", "txt"}:
            raise HTTPException(status_code=400, detail="Unsupported export type")
        if source_format not in {"markdown", "json"}:
            raise HTTPException(status_code=400, detail="Unsupported source format")

        # Generate file
        if artifact_type == "study_guide":
            if export_type == "pdf":
                file_path = generate_study_guide_pdf(
                    guide_data=req.content,
                    source_format=source_format,
                    custom_title=req.title,
                    document_title=document.title,
                )
            elif export_type == "docx":
                file_path = generate_study_guide_docx(
                    guide_data=req.content,
                    source_format=source_format,
                    custom_title=req.title,
                    document_title=document.title,
                )
            else:
                file_path = generate_study_guide_txt(
                    guide_data=req.content,
                    source_format=source_format,
                    custom_title=req.title,
                    document_title=document.title,
                )
        else:
            raise HTTPException(status_code=400, detail="Unsupported artifact type")

        # Rename to include artifact/document slug (sanitize and avoid double extensions)
        desired_ext = {"pdf": "pdf", "docx": "docx", "txt": "txt"}[export_type]
        base_name_in = req.title or f"{artifact_type}_{document.id}"
        new_filename = _safe_base_filename(base_name_in, desired_ext)
        new_path = os.path.join(EXPORTS_DIR, new_filename)
        try:
            os.replace(file_path, new_path)
            file_path = new_path
            filename = new_filename
        except Exception:
            pass

        file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0

        # Save export history (conversation_id required in schema -> use document_id placeholder)
        export_history = ExportHistory(
            user_id=current_user.id,
            conversation_id=document.id,  # placeholder to satisfy schema
            document_id=document.id,
            export_type=f"{artifact_type}_{export_type}",
            filename=filename,
            file_size=file_size,
            status="completed"
        )
        db.add(export_history)
        db.commit()
        db.refresh(export_history)

        return ExportResponse(
            success=True,
            filename=filename,
            file_size=file_size,
            download_url=f"/api/exports/download/{filename}",
            export_id=export_history.id
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Artifact export error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to export artifact: {str(e)}")

@router.get("/history", response_model=List[ExportHistoryResponse])
async def get_export_history(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Get user's export history"""
    try:
        exports = db.query(ExportHistory, Document).join(
            Document, ExportHistory.document_id == Document.id
        ).filter(
            ExportHistory.user_id == current_user.id
        ).order_by(ExportHistory.created_at.desc()).limit(50).all()
        
        return [
            ExportHistoryResponse(
                id=export_hist.id,
                document_title=doc.title,
                export_type=export_hist.export_type,
                filename=export_hist.filename,
                file_size=export_hist.file_size,
                created_at=export_hist.created_at.isoformat()
            )
            for export_hist, doc in exports
        ]
        
    except Exception as e:
        logger.error(f"Error fetching export history: {e}")
        raise HTTPException(
            status_code=500,
            detail="Failed to fetch export history"
        )

@router.delete("/history/{export_id}")
async def delete_export(
    export_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
    _=Depends(csrf_protect),
):
    """Delete an export and its file"""
    try:
        export_record = db.query(ExportHistory).filter(
            ExportHistory.id == export_id,
            ExportHistory.user_id == current_user.id
        ).first()
        
        if not export_record:
            raise HTTPException(
                status_code=404,
                detail="Export not found"
            )
        
        # Delete file if exists
        file_path = os.path.join(EXPORTS_DIR, export_record.filename)
        if os.path.exists(file_path):
            os.remove(file_path)
        
        # Delete record
        db.delete(export_record)
        db.commit()
        
        return {"message": "Export deleted successfully"}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting export: {e}")
        raise HTTPException(
            status_code=500,
            detail="Failed to delete export"
        )
